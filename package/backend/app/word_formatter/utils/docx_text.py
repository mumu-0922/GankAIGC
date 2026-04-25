"""
从 .docx 提取纯文本（用于 Requirements 预览/粗解析）。
"""
from __future__ import annotations

import io
from typing import Union

from docx import Document


def extract_text_from_docx(source: Union[str, bytes]) -> str:
    """
    从 .docx 文件提取纯文本。

    参数:
        source: 文件路径 (str) 或文件内容 (bytes)

    返回:
        提取的纯文本，段落以换行符分隔
    """
    if isinstance(source, bytes):
        doc = Document(io.BytesIO(source))
    else:
        doc = Document(source)

    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)
